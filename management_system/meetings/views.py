"""
meetings/views.py

Meeting reporting and management views.
"""

import logging
import os
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.paginator import EmptyPage, PageNotAnInteger, Paginator
from django.db.models import Q, Count
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.utils import timezone
from django.views.decorators.http import require_http_methods

from employees.models import Employee
from accounts.permissions import role_required

from .models import Meeting, MeetingNote, ActionItem, MeetingAttachment
from .forms import (
    MeetingForm, MeetingNoteForm, ActionItemForm, MeetingAttachmentForm,
    QuickMeetingForm, MeetingFilterForm
)

logger = logging.getLogger(__name__)
PAGE_SIZE = 20

# Try to import python-docx for Word document viewing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _paginate(qs, page_number, per_page=PAGE_SIZE):
    """Helper to paginate querysets."""
    paginator = Paginator(qs, per_page)
    try:
        return paginator.page(page_number)
    except PageNotAnInteger:
        return paginator.page(1)
    except EmptyPage:
        return paginator.page(paginator.num_pages)


def _extract_docx_content(file_path):
    """Extract text content from a Word document."""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document(file_path)
        content = {
            'paragraphs': [],
            'tables': [],
        }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                level = para.paragraph_format.outline_level if para.paragraph_format else 0
                content['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal',
                    'level': level,
                    'margin_left': level * 20,  # Pre-calculate margin
                })
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            if table_data:
                content['tables'].append(table_data)
        
        return content
    except Exception as e:
        logger.error(f"Error extracting DOCX content: {str(e)}")
        return None


def _get_company(request):
    """Extract company from request context."""
    return getattr(request, 'company', None)


# ---------------------------------------------------------------------------
# Meeting List & Dashboard
# ---------------------------------------------------------------------------

@role_required('admin', 'hr_manager', 'manager', 'secretary', 'accountant', 'employee')
def meeting_list(request):
    """Display all meetings with filtering and search."""
    company = _get_company(request)
    if not company:
        messages.error(request, 'No company context available.')
        return redirect('core:dashboard')
    
    meetings = Meeting.objects.filter(company=company).select_related(
        'organizer', 'created_by'
    ).prefetch_related('attendees')
    
    # Apply filters
    status = request.GET.getlist('status')
    meeting_type = request.GET.getlist('meeting_type')
    priority = request.GET.getlist('priority')
    search = request.GET.get('search', '')
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    
    if status:
        meetings = meetings.filter(status__in=status)
    if meeting_type:
        meetings = meetings.filter(meeting_type__in=meeting_type)
    if priority:
        meetings = meetings.filter(priority__in=priority)
    if search:
        meetings = meetings.filter(
            Q(title__icontains=search) |
            Q(description__icontains=search) |
            Q(location__icontains=search)
        )
    if date_from:
        meetings = meetings.filter(scheduled_date__gte=date_from)
    if date_to:
        meetings = meetings.filter(scheduled_date__lte=date_to)
    
    # Pagination
    page = request.GET.get('page', 1)
    meetings_page = _paginate(meetings.order_by('-scheduled_date'), page)
    
    context = {
        'meetings': meetings_page,
        'total_count': meetings.count(),
        'filter_form': MeetingFilterForm(request.GET or None),
    }
    
    return render(request, 'meetings/meeting_list.html', context)


@login_required
def meeting_dashboard(request):
    """Meeting dashboard with statistics and upcoming meetings."""
    company = _get_company(request)
    if not company:
        messages.error(request, 'No company context available.')
        return redirect('core:dashboard')
    
    now = timezone.now()
    
    # Statistics
    total_meetings = Meeting.objects.filter(company=company).count()
    upcoming = Meeting.objects.filter(
        company=company,
        scheduled_date__gte=now,
        status__in=['scheduled', 'in_progress']
    ).count()
    completed = Meeting.objects.filter(company=company, status='completed').count()
    pending_actions = ActionItem.objects.filter(
        meeting__company=company,
        status__in=['pending', 'in_progress']
    ).count()
    overdue_actions = ActionItem.objects.filter(
        meeting__company=company,
        status__in=['pending', 'in_progress'],
        due_date__lt=now.date()
    ).count()
    
    # Upcoming meetings
    upcoming_meetings = Meeting.objects.filter(
        company=company,
        scheduled_date__gte=now,
        status__in=['scheduled']
    ).order_by('scheduled_date')[:5].select_related('organizer').prefetch_related('attendees')
    
    # Recent meetings
    recent_meetings = Meeting.objects.filter(
        company=company
    ).order_by('-scheduled_date')[:5].select_related('organizer')
    
    # My assigned actions
    if request.user.is_authenticated:
        try:
            employee = Employee.objects.get(user=request.user, company=company)
            my_actions = ActionItem.objects.filter(
                assigned_to=employee,
                status__in=['pending', 'in_progress']
            ).order_by('due_date')[:5]
        except Employee.DoesNotExist:
            my_actions = []
    else:
        my_actions = []
    
    context = {
        'total_meetings': total_meetings,
        'upcoming_count': upcoming,
        'completed_count': completed,
        'pending_actions': pending_actions,
        'overdue_actions': overdue_actions,
        'upcoming_meetings': upcoming_meetings,
        'recent_meetings': recent_meetings,
        'my_actions': my_actions,
    }
    
    return render(request, 'meetings/meeting_dashboard.html', context)


# ---------------------------------------------------------------------------
# Meeting CRUD
# ---------------------------------------------------------------------------

@role_required('admin', 'hr_manager', 'manager', 'secretary')
def meeting_create(request):
    """Create a new meeting."""
    company = _get_company(request)
    if not company:
        messages.error(request, 'No company context available.')
        return redirect('core:dashboard')
    
    if request.method == 'POST':
        form = MeetingForm(request.POST, company=company)
        if form.is_valid():
            meeting = form.save(commit=False)
            meeting.company = company
            meeting.created_by = request.user
            meeting.save()
            form.save_m2m()  # Save many-to-many relationships
            
            messages.success(request, f'Meeting "{meeting.title}" created successfully.')
            return redirect('meetings:meeting_detail', pk=meeting.pk)
    else:
        form = MeetingForm(company=company)
    
    context = {'form': form, 'title': 'Create Meeting'}
    return render(request, 'meetings/meeting_form.html', context)


@login_required
def meeting_detail(request, pk):
    """View meeting details."""
    company = _get_company(request)
    meeting = get_object_or_404(Meeting, pk=pk, company=company)
    
    action_items = meeting.action_items.all().order_by('due_date')
    attachments = meeting.attachments.all().order_by('-uploaded_at')
    
    context = {
        'meeting': meeting,
        'action_items': action_items,
        'attachments': attachments,
        'has_notes': hasattr(meeting, 'notes'),
    }
    
    return render(request, 'meetings/meeting_detail.html', context)


@role_required('admin', 'hr_manager', 'manager', 'secretary')
def meeting_edit(request, pk):
    """Edit an existing meeting."""
    company = _get_company(request)
    meeting = get_object_or_404(Meeting, pk=pk, company=company)
    
    if request.method == 'POST':
        form = MeetingForm(request.POST, instance=meeting, company=company)
        if form.is_valid():
            form.save()
            messages.success(request, 'Meeting updated successfully.')
            return redirect('meetings:meeting_detail', pk=meeting.pk)
    else:
        form = MeetingForm(instance=meeting, company=company)
    
    context = {'form': form, 'meeting': meeting, 'title': 'Edit Meeting'}
    return render(request, 'meetings/meeting_form.html', context)


@role_required('admin', 'hr_manager', 'manager')
def meeting_delete(request, pk):
    """Delete a meeting."""
    company = _get_company(request)
    meeting = get_object_or_404(Meeting, pk=pk, company=company)
    
    if request.method == 'POST':
        title = meeting.title
        meeting.delete()
        messages.success(request, f'Meeting "{title}" deleted successfully.')
        return redirect('meetings:meeting_list')
    
    context = {'meeting': meeting}
    return render(request, 'meetings/meeting_confirm_delete.html', context)


# ---------------------------------------------------------------------------
# Meeting Notes
# ---------------------------------------------------------------------------

@role_required('admin', 'hr_manager', 'manager', 'secretary')
def meeting_notes_create(request, meeting_pk):
    """Create note for a meeting."""
    company = _get_company(request)
    meeting = get_object_or_404(Meeting, pk=meeting_pk, company=company)
    
    # Check if notes already exist
    try:
        note = meeting.notes
        return redirect('meetings:meeting_notes_edit', pk=note.pk)
    except MeetingNote.DoesNotExist:
        pass
    
    if request.method == 'POST':
        form = MeetingNoteForm(request.POST, request.FILES)
        if form.is_valid():
            note = form.save(commit=False)
            note.meeting = meeting
            note.last_edited_by = request.user
            note.save()
            messages.success(request, 'Meeting notes created successfully.')
            return redirect('meetings:meeting_detail', pk=meeting.pk)
    else:
        form = MeetingNoteForm()
    
    context = {'form': form, 'meeting': meeting, 'title': 'Create Meeting Notes'}
    return render(request, 'meetings/meeting_notes_form.html', context)


@role_required('admin', 'hr_manager', 'manager', 'secretary')
def meeting_notes_edit(request, pk):
    """Edit meeting notes."""
    note = get_object_or_404(MeetingNote, pk=pk)
    company = _get_company(request)
    
    if note.meeting.company != company:
        messages.error(request, 'Unauthorized access.')
        return redirect('core:dashboard')
    
    if request.method == 'POST':
        form = MeetingNoteForm(request.POST, request.FILES, instance=note)
        if form.is_valid():
            note = form.save(commit=False)
            note.last_edited_by = request.user
            note.save()
            messages.success(request, 'Meeting notes updated successfully.')
            return redirect('meetings:meeting_detail', pk=note.meeting.pk)
    else:
        form = MeetingNoteForm(instance=note)
    
    context = {
        'form': form,
        'meeting': note.meeting,
        'note': note,
        'title': 'Edit Meeting Notes'
    }
    return render(request, 'meetings/meeting_notes_form.html', context)


# ---------------------------------------------------------------------------
# Action Items
# ---------------------------------------------------------------------------

@role_required('admin', 'hr_manager', 'manager', 'secretary', 'employee')
def action_item_create(request, meeting_pk):
    """Create an action item for a meeting."""
    company = _get_company(request)
    meeting = get_object_or_404(Meeting, pk=meeting_pk, company=company)
    
    if request.method == 'POST':
        form = ActionItemForm(request.POST, company=company)
        if form.is_valid():
            action = form.save(commit=False)
            action.meeting = meeting
            action.save()
            messages.success(request, 'Action item created successfully.')
            return redirect('meetings:meeting_detail', pk=meeting.pk)
    else:
        form = ActionItemForm(company=company)
    
    context = {
        'form': form,
        'meeting': meeting,
        'title': 'Create Action Item'
    }
    return render(request, 'meetings/action_item_form.html', context)


@role_required('admin', 'hr_manager', 'manager', 'secretary', 'employee')
def action_item_edit(request, pk):
    """Edit an action item."""
    action_item = get_object_or_404(ActionItem, pk=pk)
    company = _get_company(request)
    
    if action_item.meeting.company != company:
        messages.error(request, 'Unauthorized access.')
        return redirect('core:dashboard')
    
    if request.method == 'POST':
        form = ActionItemForm(request.POST, instance=action_item, company=company)
        if form.is_valid():
            form.save()
            messages.success(request, 'Action item updated successfully.')
            return redirect('meetings:meeting_detail', pk=action_item.meeting.pk)
    else:
        form = ActionItemForm(instance=action_item, company=company)
    
    context = {
        'form': form,
        'action_item': action_item,
        'meeting': action_item.meeting,
        'title': 'Edit Action Item'
    }
    return render(request, 'meetings/action_item_form.html', context)


@role_required('admin', 'hr_manager', 'manager')
def action_item_delete(request, pk):
    """Delete an action item."""
    action_item = get_object_or_404(ActionItem, pk=pk)
    company = _get_company(request)
    
    if action_item.meeting.company != company:
        messages.error(request, 'Unauthorized access.')
        return redirect('core:dashboard')
    
    if request.method == 'POST':
        title = action_item.title
        meeting_pk = action_item.meeting.pk
        action_item.delete()
        messages.success(request, f'Action item "{title}" deleted successfully.')
        return redirect('meetings:meeting_detail', pk=meeting_pk)
    
    context = {'action_item': action_item}
    return render(request, 'meetings/action_item_confirm_delete.html', context)


@require_http_methods(["POST"])
@role_required('admin', 'hr_manager', 'manager', 'secretary', 'employee')
def action_item_toggle_status(request, pk):
    """Toggle action item completion status."""
    action_item = get_object_or_404(ActionItem, pk=pk)
    company = _get_company(request)
    
    if action_item.meeting.company != company:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    if action_item.status != 'completed':
        action_item.mark_completed()
        message = f'Action item marked as completed.'
    else:
        action_item.status = 'pending'
        action_item.save()
        message = f'Action item marked as pending.'
    
    messages.success(request, message)
    return redirect('meetings:meeting_detail', pk=action_item.meeting.pk)


@require_http_methods(['POST'])
@role_required('admin', 'hr_manager', 'manager', 'secretary', 'employee')
def action_item_toggle_completion(request, pk):
    """Toggle action item is_completed checkbox via AJAX."""
    import json
    action_item = get_object_or_404(ActionItem, pk=pk)
    company = _get_company(request)
    
    if action_item.meeting.company != company:
        return JsonResponse({'error': 'Unauthorized', 'success': False}, status=403)
    
    try:
        data = json.loads(request.body)
        is_completed = data.get('is_completed', False)
        
        action_item.is_completed = is_completed
        if is_completed:
            action_item.actual_completion_date = timezone.now().date()
        else:
            action_item.actual_completion_date = None
        action_item.save()
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'error': str(e), 'success': False}, status=400)


# ---------------------------------------------------------------------------
# Attachments
# ---------------------------------------------------------------------------

@role_required('admin', 'hr_manager', 'manager', 'secretary')
def attachment_upload(request, meeting_pk):
    """Upload a file attachment to a meeting."""
    company = _get_company(request)
    meeting = get_object_or_404(Meeting, pk=meeting_pk, company=company)
    
    if request.method == 'POST':
        form = MeetingAttachmentForm(request.POST, request.FILES)
        if form.is_valid():
            attachment = form.save(commit=False)
            attachment.meeting = meeting
            attachment.uploaded_by = request.user
            attachment.save()
            messages.success(request, 'File uploaded successfully.')
            return redirect('meetings:meeting_detail', pk=meeting.pk)
    else:
        form = MeetingAttachmentForm()
    
    context = {
        'form': form,
        'meeting': meeting,
        'title': 'Upload Attachment'
    }
    return render(request, 'meetings/attachment_form.html', context)


@login_required
def attachment_download(request, pk):
    """Download a meeting attachment."""
    attachment = get_object_or_404(MeetingAttachment, pk=pk)
    company = _get_company(request)
    
    if attachment.meeting.company != company:
        messages.error(request, 'Unauthorized access.')
        return redirect('core:dashboard')
    
    response = HttpResponse(attachment.file.read(), content_type='application/octet-stream')
    response['Content-Disposition'] = f'attachment; filename="{attachment.file.name}"'
    return response


@login_required
def attachment_delete(request, pk):
    """Delete an attachment."""
    attachment = get_object_or_404(MeetingAttachment, pk=pk)
    company = _get_company(request)
    
    if attachment.meeting.company != company:
        messages.error(request, 'Unauthorized access.')
        return redirect('core:dashboard')
    
    if not (request.user.is_superuser or request.user == attachment.uploaded_by):
        messages.error(request, 'You cannot delete this attachment.')
        return redirect('meetings:meeting_detail', pk=attachment.meeting.pk)
    
    if request.method == 'POST':
        meeting_pk = attachment.meeting.pk
        attachment.file.delete()
        attachment.delete()
        messages.success(request, 'Attachment deleted successfully.')
        return redirect('meetings:meeting_detail', pk=meeting_pk)
    
    context = {'attachment': attachment}
    return render(request, 'meetings/attachment_confirm_delete.html', context)


@login_required
def attachment_view(request, pk):
    """View document content (for Word files)."""
    attachment = get_object_or_404(MeetingAttachment, pk=pk)
    company = _get_company(request)
    
    if attachment.meeting.company != company:
        messages.error(request, 'Unauthorized access.')
        return redirect('core:dashboard')
    
    content = None
    is_word_doc = attachment.file.name.lower().endswith(('.doc', '.docx'))
    
    if is_word_doc and DOCX_AVAILABLE:
        try:
            file_path = attachment.file.path
            content = _extract_docx_content(file_path)
        except Exception as e:
            logger.error(f"Error viewing attachment: {str(e)}")
            messages.error(request, 'Unable to read document content.')
    elif is_word_doc and not DOCX_AVAILABLE:
        messages.warning(request, 'Word document viewing not available. Please install python-docx.')
    
    context = {
        'attachment': attachment,
        'meeting': attachment.meeting,
        'content': content,
        'is_word_doc': is_word_doc,
    }
    
    return render(request, 'meetings/attachment_view.html', context)


@login_required
def meeting_notes_view(request, pk):
    """View meeting note report document content."""
    note = get_object_or_404(MeetingNote, pk=pk)
    company = _get_company(request)
    
    if note.meeting.company != company:
        messages.error(request, 'Unauthorized access.')
        return redirect('core:dashboard')
    
    if not note.report_document:
        messages.error(request, 'No report document available.')
        return redirect('meetings:meeting_detail', pk=note.meeting.pk)
    
    content = None
    is_word_doc = note.report_document.name.lower().endswith(('.doc', '.docx'))
    
    if is_word_doc and DOCX_AVAILABLE:
        try:
            file_path = note.report_document.path
            content = _extract_docx_content(file_path)
        except Exception as e:
            logger.error(f"Error viewing notes document: {str(e)}")
            messages.error(request, 'Unable to read document content.')
    elif is_word_doc and not DOCX_AVAILABLE:
        messages.warning(request, 'Word document viewing not available. Please install python-docx.')
    
    context = {
        'note': note,
        'meeting': note.meeting,
        'content': content,
        'is_word_doc': is_word_doc,
        'document_name': note.report_document.name.split('/')[-1],
    }
    
    return render(request, 'meetings/meeting_notes_view.html', context)


# ---------------------------------------------------------------------------
# Reports & Analytics
# ---------------------------------------------------------------------------

@role_required('admin', 'hr_manager', 'manager', 'accountant')
def meeting_report(request):
    """Generate meeting analytics and reports."""
    company = _get_company(request)
    if not company:
        messages.error(request, 'No company context available.')
        return redirect('core:dashboard')
    
    # Meeting statistics
    meetings = Meeting.objects.filter(company=company)
    total_meetings = meetings.count()
    meetings_by_type = dict(meetings.values_list('meeting_type').annotate(Count('id')))
    meetings_by_status = dict(meetings.values_list('status').annotate(Count('id')))
    
    # Action items statistics
    actions = ActionItem.objects.filter(meeting__company=company)
    total_actions = actions.count()
    completed_actions = actions.filter(status='completed').count()
    overdue_actions = actions.filter(status__in=['pending', 'in_progress'], due_date__lt=timezone.now().date()).count()
    
    # Attendance statistics
    top_attendees = []
    
    # Calculate rates
    completion_rate = (completed_actions / total_actions * 100) if total_actions > 0 else 0
    overdue_rate = (overdue_actions / total_actions * 100) if total_actions > 0 else 0
    
    context = {
        'total_meetings': total_meetings,
        'meetings_by_type': meetings_by_type,
        'meetings_by_status': meetings_by_status,
        'total_actions': total_actions,
        'completed_actions': completed_actions,
        'completion_rate': completion_rate,
        'overdue_actions': overdue_actions,
        'overdue_rate': overdue_rate,
    }
    
    return render(request, 'meetings/meeting_report.html', context)
